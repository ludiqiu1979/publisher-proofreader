from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import List, Tuple
import re
from datetime import datetime
import uuid
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Word 文档处理器"""
    
    def __init__(self, doc_path: str):
        """
        初始化文档处理器
        
        Args:
            doc_path: Word 文档路径
        """
        self.doc_path = Path(doc_path)
        self.doc = Document(doc_path)
        self.revisions_enabled = False
        self.author = "Proofreader"
        self.revision_id = 0
        
    def enable_track_changes(self, author: str = "Proofreader"):
        """启用修订跟踪"""
        self.author = author
        settings = self.doc.settings.element
        track_revs = settings.find(qn('w:trackRevisions'))
        if track_revs is None:
            track_revs = OxmlElement('w:trackRevisions')
            settings.append(track_revs)
        self.revisions_enabled = True
        logger.info(f"Track Changes 已启用, 作者: {author}")
    
    def get_paragraphs_text(self) -> List[str]:
        """获取所有段落文本"""
        return [para.text for para in self.doc.paragraphs]
    
    def get_paragraph_by_index(self, index: int):
        """根据索引获取段落"""
        if 0 <= index < len(self.doc.paragraphs):
            return self.doc.paragraphs[index]
        return None
    
    def count_words(self) -> int:
        """计算单词数"""
        total_words = 0
        for para in self.doc.paragraphs:
            words = para.text.split()
            total_words += len(words)
        return total_words
    
    def count_sentences(self) -> int:
        """计算句数"""
        total_text = "\n".join([p.text for p in self.doc.paragraphs])
        sentences = re.split(r'[。！？\.\!\?]+', total_text)
        return len([s for s in sentences if s.strip()])
    
    def add_tracked_deletion(self, paragraph_index: int, text_to_delete: str, 
                            author: str = None, comment: str = None) -> bool:
        """添加删除修订"""
        author = author or self.author
        para = self.get_paragraph_by_index(paragraph_index)
        
        if not para:
            logger.error(f"段落索引 {paragraph_index} 不存在")
            return False
        
        try:
            W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            p_element = para._element
            
            del_elem = OxmlElement('w:del')
            del_elem.set(f'{W}id', str(self.revision_id))
            del_elem.set(f'{W}author', author)
            del_elem.set(f'{W}date', datetime.utcnow().isoformat() + 'Z')
            
            run_elem = OxmlElement('w:r')
            del_elem.append(run_elem)
            
            del_text_elem = OxmlElement('w:delText')
            del_text_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            del_text_elem.text = text_to_delete
            run_elem.append(del_text_elem)
            
            p_element.append(del_elem)
            self.revision_id += 1
            
            logger.info(f"已添加删除修订: {text_to_delete[:50]}")
            return True
            
        except Exception as e:
            logger.error(f"添加删除修订失败: {e}")
            return False
    
    def add_tracked_insertion(self, paragraph_index: int, text_to_insert: str,
                             author: str = None, comment: str = None) -> bool:
        """添加插入修订"""
        author = author or self.author
        para = self.get_paragraph_by_index(paragraph_index)
        
        if not para:
            logger.error(f"段落索引 {paragraph_index} 不存在")
            return False
        
        try:
            W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            p_element = para._element
            
            ins_elem = OxmlElement('w:ins')
            ins_elem.set(f'{W}id', str(self.revision_id))
            ins_elem.set(f'{W}author', author)
            ins_elem.set(f'{W}date', datetime.utcnow().isoformat() + 'Z')
            
            run_elem = OxmlElement('w:r')
            ins_elem.append(run_elem)
            
            text_elem = OxmlElement('w:t')
            text_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            text_elem.text = text_to_insert
            run_elem.append(text_elem)
            
            p_element.append(ins_elem)
            self.revision_id += 1
            
            logger.info(f"已添加插入修订: {text_to_insert[:50]}")
            return True
            
        except Exception as e:
            logger.error(f"添加插入修订失败: {e}")
            return False
    
    def add_comment(self, paragraph_index: int, comment_text: str, 
                   author: str = None) -> bool:
        """添加评论"""
        author = author or self.author
        para = self.get_paragraph_by_index(paragraph_index)
        
        if not para:
            return False
        
        try:
            para.add_run(f"\n[{author}的批注: {comment_text}]").italic = True
            return True
        except Exception as e:
            logger.error(f"添加评论失败: {e}")
            return False
    
    def save_document(self, output_path: str) -> bool:
        """保存文档"""
        try:
            self.doc.save(output_path)
            logger.info(f"文档已保存到: {output_path}")
            return True
        except Exception as e:
            logger.error(f"保存文档失败: {e}")
            return False
    
    def get_document_stats(self) -> dict:
        """获取文档统计信息"""
        return {
            "total_paragraphs": len(self.doc.paragraphs),
            "total_words": self.count_words(),
            "total_sentences": self.count_sentences(),
            "total_characters": sum(len(p.text) for p in self.doc.paragraphs),
        }
